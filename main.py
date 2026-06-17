from fastapi import FastAPI, Request
from fastapi.responses import FileResponse, StreamingResponse
import json
import requests
from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml import OxmlElement
from io import BytesIO
import urllib.parse
from PIL import Image  # ✅ for dynamic sizing

app = FastAPI()

SUPABASE_URL = "https://gdcwjpkgffqmatsmuqra.supabase.co"
BASE_URL = f"{SUPABASE_URL}/storage/v1/object/public/question-images"


# ✅ HOME
@app.get("/")
def serve_home():
    return FileResponse("index.html")


# ✅ STRUCTURE
@app.get("/structure")
def get_structure():
    with open("structure.json") as f:
        return json.load(f)


# ✅ GENERATE
@app.post("/generate")
async def generate(request: Request):

    data = await request.json()
    entries = data["entries"]
    filetype = data["filetype"]

    doc = Document()
    pdf_data = []

    for paper_name, q in entries:

        parts = paper_name.split(" ")
        month = parts[0]

        if month == "November":
            month = "Nov"

        year = parts[1][-2:]
        paper = parts[2]

        paper_fixed = f"{month} {year} {paper}"
        filename_base = f"{paper_fixed}_Q{q}"

        images = []

        base_url = f"{BASE_URL}/{urllib.parse.quote(filename_base + '.png')}"
        res = requests.get(base_url)

        if res.status_code == 200:
            images.append(res.content)
        else:
            for i in range(1, 6):
                part_url = f"{BASE_URL}/{urllib.parse.quote(filename_base + '_' + str(i) + '.png')}"
                res = requests.get(part_url)
                if res.status_code == 200:
                    images.append(res.content)

        # ✅ MISSING
        if not images:
            if filetype == "word":
                doc.add_paragraph(f"MISSING: {paper_fixed} Q{q}")
            continue

        # =================
        # ✅ WORD (FINAL VERSION)
        # =================
        if filetype == "word":

            table = doc.add_table(rows=1, cols=1)
            row = table.rows[0]
            cell = row.cells[0]

            # ✅ Prevent split across pages
            tr = row._tr
            trPr = tr.get_or_add_trPr()
            cant_split = OxmlElement('w:cantSplit')
            trPr.append(cant_split)

            # ✅ Header
            p = cell.paragraphs[0]
            run = p.add_run(f"{paper_fixed} Question {q}")
            p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            run.bold = True

            # ✅ Dynamic images
            for img in images:
                paragraph = cell.add_paragraph()
                run = paragraph.add_run()

                image_stream = BytesIO(img)
                pil_img = Image.open(image_stream)

                orig_width, orig_height = pil_img.size
                dpi = pil_img.info.get("dpi", (96, 96))[0]

                width_in_inches = orig_width / dpi

                max_width_inches = 6.5  # ✅ page-fit width
                scale = min(1, max_width_inches / width_in_inches)

                final_width = width_in_inches * scale

                image_stream.seek(0)

                run.add_picture(image_stream, width=Inches(final_width))

            doc.add_paragraph("")

        # =================
        # ✅ STORE FOR PDF
        # =================
        if filetype == "pdf":
            pdf_data.append((paper_fixed, q, images))

    # =================
    # ✅ WORD OUTPUT
    # =================
    if filetype == "word":

        stream = BytesIO()
        doc.save(stream)
        stream.seek(0)

        return StreamingResponse(
            stream,
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            headers={"Content-Disposition": "attachment; filename=generated.docx"}
        )

    # =================
    # ✅ PDF OUTPUT (FINAL)
    # =================
    if filetype == "pdf":

        from reportlab.pdfgen import canvas
        from reportlab.lib.pagesizes import A4
        from reportlab.lib.utils import ImageReader

        pdf_stream = BytesIO()
        c = canvas.Canvas(pdf_stream, pagesize=A4)

        PAGE_WIDTH, PAGE_HEIGHT = A4
        y = PAGE_HEIGHT - 40

        for paper_fixed, q, images in pdf_data:

            header = f"{paper_fixed} Question {q}"
            header_height = 30

            total_height = header_height
            scaled = []

            for img_bytes in images:
                img_reader = ImageReader(BytesIO(img_bytes))
                w, h = img_reader.getSize()

                scale = (PAGE_WIDTH - 80) / w
                new_h = h * scale

                total_height += new_h + 20
                scaled.append((img_reader, new_h))

            total_height += 20

            # ✅ Move whole question if needed
            if y - total_height < 50:
                c.showPage()
                y = PAGE_HEIGHT - 40

            # ✅ Header
            c.setFont("Helvetica-Bold", 14)
            c.drawCentredString(PAGE_WIDTH / 2, y, header)
            y -= header_height

            # ✅ Images
            for img_reader, img_height in scaled:
                c.drawImage(
                    img_reader,
                    40,
                    y - img_height,
                    width=PAGE_WIDTH - 80,
                    height=img_height
                )
                y -= (img_height + 20)

            y -= 20

        c.save()
        pdf_stream.seek(0)

        return StreamingResponse(
            pdf_stream,
            media_type="application/pdf",
            headers={"Content-Disposition": "attachment; filename=generated.pdf"}
        )
